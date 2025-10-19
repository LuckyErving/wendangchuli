#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""测试Word转PDF的中文支持"""

import os
import sys

# 测试文件夹
test_folder = "/Users/ervin/Downloads/630222203210020071(李时海)"

if not os.path.exists(test_folder):
    print(f"错误：测试文件夹不存在: {test_folder}")
    sys.exit(1)

# 查找docx文件
docx_files = []
for root, dirs, filenames in os.walk(test_folder):
    for filename in filenames:
        if filename.lower().endswith('.docx') and not filename.startswith('~'):
            filepath = os.path.join(root, filename)
            docx_files.append(filepath)

if not docx_files:
    print("没有找到.docx文件")
    sys.exit(0)

print(f"找到 {len(docx_files)} 个.docx文件:")
for f in docx_files:
    print(f"  - {os.path.basename(f)}")

print("\n选择第一个文件进行测试...")
test_file = docx_files[0]
print(f"测试文件: {test_file}")

# 测试转换
from docx import Document
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.enums import TA_LEFT, TA_CENTER

# 读取Word文档
print("\n读取Word文档...")
doc = Document(test_file)
print(f"  找到 {len(doc.paragraphs)} 个段落")

# 显示前3个段落
print("\n前3个段落内容:")
for i, para in enumerate(doc.paragraphs[:3]):
    if para.text.strip():
        print(f"  {i+1}. {para.text[:50]}...")

# 创建测试PDF
output_pdf = os.path.join(os.path.dirname(test_file), "test_中文转换.pdf")
print(f"\n生成PDF: {output_pdf}")

pdf_doc = SimpleDocTemplate(output_pdf, pagesize=A4)
story = []

# 注册中文字体
pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
print("  ✓ 中文字体已注册")

# 创建样式
styles = getSampleStyleSheet()
chinese_style = ParagraphStyle(
    'Chinese',
    parent=styles['Normal'],
    fontName='STSong-Light',
    fontSize=12,
    leading=20,
    alignment=TA_LEFT,
)

title_style = ParagraphStyle(
    'ChineseTitle',
    parent=styles['Heading1'],
    fontName='STSong-Light',
    fontSize=16,
    leading=24,
    alignment=TA_CENTER,
)

# 添加标题
filename = os.path.basename(test_file)
story.append(Paragraph(filename, title_style))
story.append(Spacer(1, 20))

# 添加段落
para_count = 0
for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        text = paragraph.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        try:
            p = Paragraph(text, chinese_style)
            story.append(p)
            story.append(Spacer(1, 12))
            para_count += 1
        except Exception as e:
            print(f"  ⚠️  段落处理失败: {str(e)}")

print(f"  添加了 {para_count} 个段落")

# 生成PDF
print("\n生成PDF文件...")
pdf_doc.build(story)
print(f"✓ PDF生成成功: {output_pdf}")
print(f"\n文件大小: {os.path.getsize(output_pdf)} 字节")

print("\n测试完成！请打开PDF文件查看中文显示效果。")
